from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt


def export_markdown(text: str, out_path: Path) -> Path:
    out_path.write_text(text, encoding="utf-8")
    return out_path


def export_text(text: str, out_path: Path) -> Path:
    plain = re.sub(r"[*_`#>|]", "", text)
    out_path.write_text(plain, encoding="utf-8")
    return out_path


def export_docx(markdown_text: str, out_path: Path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Tahoma"
    style.font.size = Pt(11)
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|$", lines[i+1].strip()):
            table_lines = [line]
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            _add_table(doc, [line] + table_lines[1:])
            i = j
            continue
        if re.match(r"^\s*[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", line), style="List Bullet")
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        doc.add_paragraph(line)
        i += 1
    doc.save(out_path)
    return out_path


def _add_table(doc, table_lines):
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in table_lines if ln.strip()]
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid"
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            if c < len(table.rows[r].cells):
                table.rows[r].cells[c].text = val
